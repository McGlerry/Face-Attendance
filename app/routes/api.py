# app/routes/api.py
from flask import Blueprint, request, jsonify, Response, session, current_app
import os
import csv
import threading
import time
import logging
from datetime import datetime
import cv2
import numpy as np
from functools import wraps

logger = logging.getLogger(__name__)

from .. import config_manager, db_manager, attendance_system

# Import training functions
from ..services.training import enhanced_train_model_cnn_optimized

# Global training thread reference
training_thread = None

# Constants
EXPORTS_DIR = "data/exports"
REPORTLAB_AVAILABLE = True
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# DOCX availability check
DOCX_AVAILABLE = False
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

api_bp = Blueprint('api', __name__, url_prefix='/api')

# Authentication decorator
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return jsonify({'success': False, 'message': 'Authentication required', 'redirect': '/login'}), 401
        return f(*args, **kwargs)
    return decorated_function

# Utility functions
def _safe_name(name):
    """Create a safe filename from a name"""
    return "".join(c for c in name if c.isalnum() or c in (' ', '-', '_')).rstrip()

@api_bp.route('/capture_video_feed')
def capture_video_feed():
    def generate_capture_frame():
        while True:
            frame = attendance_system.camera_manager.get_frame()
            if frame is None:
                time.sleep(0.05)
                continue

            # Overlay for capture feed (bounding box, quality message)
            rgb_frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            detection_model = config_manager.config['recognition_settings'].get(
                'capture_detection_model', 'hog')
            face_locations = attendance_system.face_model.detect_faces(
                rgb_frame, model=detection_model)

            if config_manager.config['system_settings']['enable_quality_validation'] and face_locations:
                # Get the largest face
                face_areas = [(right-left)*(bottom-top)
                              for (top, right, bottom, left) in face_locations]
                idx = int(np.argmax(face_areas))
                face_location = face_locations[idx]
                top, right, bottom, left = face_location

                quality_ok, quality_msg = attendance_system.face_model.validate_classroom_face_quality(
                    face_location, frame
                )

                color = (0, 255, 0) if quality_ok else (
                    0, 0, 255)  # Green for good, Red for bad
                cv2.rectangle(frame, (left, top), (right, bottom), color, 2)
                cv2.putText(frame, quality_msg, (left, top - 10),
                            cv2.FONT_HERSHEY_SIMPLEX, 0.7, color, 2)

            ret, buffer = cv2.imencode('.jpg', frame, [
                                       cv2.IMWRITE_JPEG_QUALITY, config_manager.config['display_settings']['jpeg_quality']])
            if not ret:
                continue
            yield (b'--frame\r\n'
                   b'Content-Type: image/jpeg\r\n\r\n' + buffer.tobytes() + b'\r\n')
    return Response(generate_capture_frame(), mimetype='multipart/x-mixed-replace; boundary=frame')

@api_bp.route('/start_capture_camera', methods=['POST'])
def start_capture_camera():
    success = attendance_system.camera_manager.start_camera()
    return jsonify({'success': success})

@api_bp.route('/stop_capture_camera', methods=['POST'])
def stop_capture_camera():
    attendance_system.camera_manager.stop_camera()
    return jsonify({'success': True})

@api_bp.route('/capture_face/<int:student_id>', methods=['POST'])
@login_required
def capture_face(student_id):
    """Enhanced face capture with quality validation"""
    student = db_manager.get_student_by_id(student_id)
    if not student:
        return jsonify({'success': False, 'message': 'Student not found.'})

    try:
        # Check if image data is provided
        if 'image' not in request.files:
            return jsonify({'success': False, 'message': 'No image data provided'})

        image_file = request.files['image']
        if image_file.filename == '':
            return jsonify({'success': False, 'message': 'No image selected'})

        # Read image data
        image_data = image_file.read()
        nparr = np.frombuffer(image_data, np.uint8)
        frame = cv2.imdecode(nparr, cv2.IMREAD_COLOR)

        if frame is None:
            return jsonify({'success': False, 'message': 'Invalid image format'})

        # Process the frame for face detection (using HOG for faster capture performance)
        rgb_frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
        detection_model = config_manager.config['recognition_settings'].get(
            'capture_detection_model', 'hog')
        face_locations = attendance_system.face_model.detect_faces(
            rgb_frame, model=detection_model)

        if not face_locations:
            return jsonify({'success': False, 'message': 'No face detected. Please position your face clearly in the camera.'})

        # Get the largest face
        face_areas = [(right-left)*(bottom-top)
                      for (top, right, bottom, left) in face_locations]
        idx = int(np.argmax(face_areas))
        face_location = face_locations[idx]
        top, right, bottom, left = face_location

        # Prepare face_location for return
        face_location_list = list(face_location)

        # ENHANCED: Quality validation
        quality_ok, quality_msg = True, "N/A"  # Default if disabled
        if config_manager.config['system_settings']['enable_quality_validation']:
            quality_ok, quality_msg = attendance_system.face_model.validate_classroom_face_quality(
                face_location, frame
            )
            if not quality_ok:
                return jsonify({
                    'success': False,
                    'message': f'Image quality issue: {quality_msg}. Please adjust lighting, position, or angle.',
                    'face_location': face_location_list,
                    'quality_message': quality_msg
                })

        # Extract face with padding
        padding = 20
        y1 = max(0, top-padding)
        y2 = min(frame.shape[0], bottom+padding)
        x1 = max(0, left-padding)
        x2 = min(frame.shape[1], right-padding)
        face_roi = frame[y1:y2, x1:x2]

        if face_roi.size == 0:
            return jsonify({'success': False, 'message': 'Invalid face region detected'})

        # Resize and optimize image storage
        face_roi = cv2.resize(face_roi, (150, 150))
        face_roi = attendance_system.memory_optimizer.optimize_image_storage(
            face_roi)
        safe_name = _safe_name(student['name'])

        # Determine next available index
        dataset_dir = config_manager.config['system_settings']['dataset_dir']
        existing = [f for f in os.listdir(
            dataset_dir) if f.startswith(f"{student['id']}_")]
        next_idx = 1
        if existing:
            nums = []
            for f in existing:
                try:
                    parts = f.split('_')
                    if len(parts) >= 3:
                        num = int(parts[-1].split('.')[0])
                        nums.append(num)
                except Exception:
                    pass
            if nums:
                next_idx = max(nums) + 1

        filename = os.path.join(
            dataset_dir, f"{student['id']}_{safe_name}_{next_idx:03d}.jpg")
        cv2.imwrite(filename, face_roi, [cv2.IMWRITE_JPEG_QUALITY, 95])

        # Get current capture count for this student
        current_captures = len([f for f in os.listdir(
            dataset_dir) if f.startswith(f"{student_id}_")])
        max_captures = config_manager.config['recognition_settings']['capture_images_per_student']

        # Log capture event
        db_manager.log_security_event("FACE_CAPTURED", student['id'],
                                      details=f"Image {current_captures}/{max_captures}")

        completed = current_captures >= max_captures
        if completed:
            db_manager.log_security_event("CAPTURE_COMPLETED", student['id'],
                                          details=f"Total images: {current_captures}")

        return jsonify({
            'success': True,
            'message': f'High-quality face captured successfully ({current_captures}/{max_captures})',
            'progress': (current_captures / max_captures) * 100,
            'count': current_captures,
            'total': max_captures,
            'completed': completed,
            'quality_validated': quality_ok,
            'quality_message': quality_msg,
            'face_location': face_location_list
        })

    except (ValueError, KeyError, sqlite3.Error) as e:
        logger.exception(f"Capture face error for student {student_id}: {e}")
        return jsonify({'success': False, 'message': f'Invalid request: {str(e)}'})
    except Exception as e:
        logger.exception(f"Unexpected capture face error: {e}")
        return jsonify({'success': False, 'message': 'Server error during capture'})

@api_bp.route('/start_system', methods=['POST'])
@login_required
def start_system():
    try:
        success, message = attendance_system.start_system()
        return jsonify({'success': success, 'message': message})
    except (RuntimeError, ValueError) as e:
        logger.exception(f"System start error: {e}")
        return jsonify({'success': False, 'message': f'System error: {str(e)}'})
    except Exception as e:
        logger.exception(f"Unexpected system start error: {e}")
        return jsonify({'success': False, 'message': 'Server error starting system'})

@api_bp.route('/stop_system', methods=['POST'])
@login_required
def stop_system():
    try:
        attendance_system.stop_system()
        return jsonify({'success': True, 'message': 'System stopped successfully'})
    except RuntimeError as e:
        logger.exception(f"System stop error: {e}")
        return jsonify({'success': False, 'message': f'System error: {str(e)}'})
    except Exception as e:
        logger.exception(f"Unexpected system stop error: {e}")
        return jsonify({'success': False, 'message': 'Server error stopping system'})

@api_bp.route('/system_status')
def system_status():
    try:
        status = attendance_system.get_system_status()
        return jsonify(status)
    except Exception as e:
        return jsonify({
            'running': False,
            'camera_active': False,
            'error': str(e),
            'stats': {
                'faces_detected': 0, 'recognitions': 0, 'attendances_marked': 0, 'camera_fps': 0.0,
                'faces_rejected_quality': 0, 'auto_mark_matches': 0, 'high_confidence_matches': 0,
                'standard_matches': 0, 'suspicious_matches': 0, 'processing_fps': 0.0,
                'average_processing_time_ms': 0.0
            },
            'active_classes': [],
            'model_loaded': False,
            'recognition_buffer_size': 0,
            'security_features': {
                'quality_validation': config_manager.config['system_settings']['enable_quality_validation'],
                'multiple_confirmations': config_manager.config['system_settings']['enable_multiple_confirmations'],
                'security_logging': config_manager.config['system_settings']['security_logging']
            }
        })

@api_bp.route('/train_model', methods=['POST'])
@login_required
def train_model_route():
    def run_training():
        try:
            # Always call CNN training
            enhanced_train_model_cnn_optimized(
                config_manager, db_manager, attendance_system)
        except Exception as e:
            logger.exception(f"Training failed unexpectedly: {e}")

    global training_thread
    training_thread = threading.Thread(target=run_training, daemon=True)
    training_thread.start()

    return jsonify({'success': True, 'message': 'CNN model training started. This will take longer but provide better accuracy.'})

@api_bp.route('/training_progress')
def get_training_progress_route():
    from ..services.training import get_training_progress
    return jsonify(get_training_progress())

@api_bp.route('/generate_report', methods=['POST'])
def generate_report():
    try:
        start_date = request.form['start_date']
        end_date = request.form['end_date']
        class_id = request.form.get('class_id')

        # Validate dates
        try:
            datetime.strptime(start_date, '%Y-%m-%d')
            datetime.strptime(end_date, '%Y-%m-%d')
        except ValueError:
            return jsonify({
                'success': False,
                'error': 'Invalid date format. Use YYYY-MM-DD'
            })

        if class_id and class_id.lower() != 'all':
            try:
                class_id = int(class_id)
            except ValueError:
                return jsonify({
                    'success': False,
                    'error': 'Invalid class ID'
                })
        else:
            class_id = None

        records = db_manager.get_attendance_records(
            start_date, end_date, class_id)

        # Format records for frontend - use consistent key naming
        formatted_records = []
        for record in records:
            formatted_records.append({
                'student_id': record['student_id'],
                'name': record['name'],
                'class_name': record['class_name'],
                'timestamp': record['timestamp'],
                'confidence_score': record['confidence_score'] if record['confidence_score'] is not None else 0.0,
                'status': record['status'] if record['status'] else 'present',
                'match_type': record['match_type'] if record['match_type'] else 'standard',
                'date': record['date'],
                'class_id': record['class_id']
            })

        return jsonify({
            'success': True,
            'records': formatted_records,
            'count': len(formatted_records),
            'date_range': f"{start_date} to {end_date}",
            'class_filter': class_id
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'Failed to generate report: {str(e)}'
        })

@api_bp.route('/export_csv', methods=['POST'])
def export_csv():
    try:
        start_date = request.form['start_date']
        end_date = request.form['end_date']
        class_id = request.form.get('class_id')

        records = db_manager.get_attendance_records(
            start_date, end_date, class_id)

        if not records:
            return jsonify({'success': False, 'message': 'No records found for export.'})

        # Group by class
        from collections import defaultdict
        import io

        grouped_by_class = defaultdict(list)
        for record in records:
            grouped_by_class[record['class_name']].append(record)

        # Generate CSV with class sections
        output = io.StringIO()
        writer = csv.writer(output)

        # Write each class as a separate section
        for class_name in sorted(grouped_by_class.keys()):
            class_records = grouped_by_class[class_name]
            
            # Class header
            writer.writerow([])
            writer.writerow([f"Class: {class_name}"])
            writer.writerow([f"Total Records: {len(class_records)}"])
            writer.writerow([])
            
            # Table headers
            writer.writerow(['Student ID', 'Student Name', 'Date', 'Timestamp', 'Status', 'Confidence Score', 'Match Type'])
            
            # Sort records by student name and date
            class_records.sort(key=lambda x: (x['name'], x['date']))
            
            # Write records
            for record in class_records:
                timestamp = record['timestamp'] if record['timestamp'] is not None else record['date']
                confidence = f"{record['confidence_score']:.3f}" if record['confidence_score'] is not None else '0.000'
                status = record['status'].title() if record['status'] else 'Present'
                match_type = record['match_type'] if record['match_type'] else 'Standard'
                
                writer.writerow([
                    record['student_id'],
                    record['name'],
                    record['date'],
                    timestamp,
                    status,
                    confidence,
                    match_type
                ])
            
            writer.writerow([])  # Empty row between classes

        csv_content = output.getvalue()
        output.close()

        # Define filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"attendance_by_class_{timestamp}.csv"

        db_manager.log_security_event(
            "REPORT_EXPORTED", details=f"CSV report by class generated: {filename}")
        return Response(
            csv_content,
            mimetype='text/csv',
            headers={'Content-Disposition': f'attachment; filename={filename}'}
        )

    except Exception as e:
        return jsonify({'success': False, 'message': f'Failed to export CSV: {str(e)}'})

@api_bp.route('/export_pdf', methods=['POST'])
def export_pdf():
    try:
        if not REPORTLAB_AVAILABLE:
            return jsonify({'success': False, 'message': 'PDF export is not available. ReportLab library is not installed.'})

        start_date = request.form['start_date']
        end_date = request.form['end_date']
        class_id = request.form.get('class_id')

        records = db_manager.get_attendance_records(
            start_date, end_date, class_id)

        if not records:
            return jsonify({'success': False, 'message': 'No records found for export.'})

        # Group records by class and date for better organization
        from collections import defaultdict
        grouped_data = defaultdict(lambda: defaultdict(list))

        for record in records:
            class_name = record['class_name']
            date = record['date']
            grouped_data[class_name][date].append(record)

        # Generate PDF
        from io import BytesIO
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        elements = []

        # Title
        title = Paragraph(
            f"Attendance Report ({start_date} to {end_date})", getSampleStyleSheet()['Title'])
        elements.append(title)
        elements.append(Spacer(1, 12))

        # Summary
        total_records = len(records)
        present_count = sum(1 for r in records if r['status'] in ['present', 'late'])
        absent_count = sum(1 for r in records if r['status'] == 'absent')
        
        summary_text = f"Total Records: {total_records} | Present/Late: {present_count} | Absent: {absent_count}"
        summary = Paragraph(summary_text, getSampleStyleSheet()['Normal'])
        elements.append(summary)
        elements.append(Spacer(1, 12))

        # Group by class
        for class_name in sorted(grouped_data.keys()):
            # Class header
            class_title = Paragraph(f"Class: {class_name}", getSampleStyleSheet()['Heading2'])
            elements.append(class_title)
            elements.append(Spacer(1, 6))

            class_dates = grouped_data[class_name]
            
            for date in sorted(class_dates.keys()):
                # Date header
                date_title = Paragraph(f"Date: {date}", getSampleStyleSheet()['Heading3'])
                elements.append(date_title)
                elements.append(Spacer(1, 6))

                # Table for this date
                date_records = class_dates[date]
                data = [['Student ID', 'Student Name', 'Status', 'Timestamp', 'Confidence']]
                
                for record in sorted(date_records, key=lambda x: x['name']):
                    timestamp = record['timestamp'] if record['timestamp'] else 'N/A'
                    confidence = f"{record['confidence_score']:.3f}" if record['confidence_score'] else '0.000'
                    data.append([
                        str(record['student_id']),
                        record['name'],
                        record['status'].title() if record['status'] else 'Present',
                        timestamp,
                        confidence
                    ])

                if len(data) > 1:  # Only create table if there are records
                    table = Table(data)
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('FONTSIZE', (0, 1), (-1, -1), 8),
                    ]))
                    elements.append(table)
                    elements.append(Spacer(1, 12))

        # Build PDF
        doc.build(elements)
        pdf_content = buffer.getvalue()
        buffer.close()

        # Define filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"attendance_by_class_{timestamp}.pdf"

        db_manager.log_security_event(
            "REPORT_EXPORTED", details=f"PDF report by class generated: {filename}")
        return Response(
            pdf_content,
            mimetype='application/pdf',
            headers={'Content-Disposition': f'attachment; filename={filename}'}
        )

    except Exception as e:
        return jsonify({'success': False, 'message': f'Failed to export PDF: {str(e)}'})

@api_bp.route('/export_docx', methods=['POST'])
def export_docx():
    try:
        if not DOCX_AVAILABLE:
            return jsonify({'success': False, 'message': 'Word export is not available. python-docx library is not installed.'})

        start_date = request.form['start_date']
        end_date = request.form['end_date']
        class_id = request.form.get('class_id')

        records = db_manager.get_attendance_records(
            start_date, end_date, class_id)

        if not records:
            return jsonify({'success': False, 'message': 'No records found for export.'})

        # Group records by class and date (match PDF/CSV structure)
        from collections import defaultdict
        grouped_data = defaultdict(lambda: defaultdict(list))
        for record in records:
            class_name = record['class_name']
            date = record['date']
            grouped_data[class_name][date].append(record)

        # Create DOCX document
        from io import BytesIO
        buffer = BytesIO()
        doc = Document()

        # Title
        title = doc.add_heading(f'Attendance Report ({start_date} to {end_date})', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Summary statistics (match PDF)
        total_records = len(records)
        present_count = sum(1 for r in records if r['status'] in ['present', 'late'])
        absent_count = sum(1 for r in records if r['status'] == 'absent')
        late_count = sum(1 for r in records if r['status'] == 'late')

        summary_p = doc.add_paragraph()
        summary_p.add_run(f'Total Records: {total_records} | Present: {present_count - late_count} | Late: {late_count} | Absent: {absent_count}').bold = True
        summary_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        # Group by class (sorted)
        for class_name in sorted(grouped_data.keys()):
            # Class heading
            class_heading = doc.add_heading(f'Class: {class_name}', level=1)

            class_dates = grouped_data[class_name]
            for date in sorted(class_dates.keys()):
                # Date sub-heading
                date_heading = doc.add_heading(f'Date: {date}', level=2)

                # Table for this date/class
                date_records = sorted(class_dates[date], key=lambda x: x['name'])
                if date_records:
                    table = doc.add_table(rows=1, cols=6)
                    table.style = 'Table Grid'
                    table.autofit = True

                    # Header row
                    hdr_cells = table.rows[0].cells
                    headers = ['Student ID', 'Student Name', 'Status', 'Timestamp', 'Confidence', 'Match Type']
                    for i, header in enumerate(headers):
                        hdr_cells[i].text = header
                        hdr_cells[i].paragraphs[0].runs[0].bold = True

                    # Data rows
                    for record in date_records:
                        row_cells = table.add_row().cells
                        timestamp = record['timestamp'] if record['timestamp'] else record['date']
                        confidence = f"{record['confidence_score']:.3f}" if record['confidence_score'] else '0.000'
                        status = record['status'].title() if record['status'] else 'Present'
                        match_type = record['match_type'] if record['match_type'] else 'Standard'

                        row_cells[0].text = str(record['student_id'])
                        row_cells[1].text = record['name']
                        row_cells[2].text = status
                        row_cells[3].text = timestamp
                        row_cells[4].text = confidence
                        row_cells[5].text = match_type

        # Footer
        footer_p = doc.add_paragraph('Report generated on ' + datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Save to buffer
        doc.save(buffer)
        docx_content = buffer.getvalue()
        buffer.close()

        # Filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"attendance_by_class_{timestamp}.docx"

        db_manager.log_security_event(
            "REPORT_EXPORTED", details=f"Word report by class generated: {filename}")

        return Response(
            docx_content,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': f'attachment; filename={filename}'}
        )

    except Exception as e:
        logger.exception("DOCX export error")
        return jsonify({'success': False, 'message': f'Failed to export Word: {str(e)}'})








@api_bp.route('/restart_camera', methods=['POST'])
def restart_camera():
    """Restart camera without restarting full system"""
    try:
        attendance_system.camera_manager.stop_camera()
        time.sleep(1)

        success = attendance_system.camera_manager.start_camera()

        if success:
            db_manager.log_security_event(
                "CAMERA_RESTARTED", details="Camera manually restarted")
            return jsonify({'success': True, 'message': 'Camera restarted successfully'})
        else:
            return jsonify({'success': False, 'message': 'Failed to restart camera. Check camera connection.'})
    except Exception as e:
        return jsonify({'success': False, 'message': f'Error restarting camera: {str(e)}'})

@api_bp.route('/start_camera_only', methods=['POST'])
def start_camera_only():
    """Start camera in monitor mode (no recognition)"""
    try:
        if attendance_system.camera_manager.is_running:
            return jsonify({'success': True, 'message': 'Camera already running'})

        success = attendance_system.camera_manager.start_camera()

        if success:
            db_manager.log_security_event(
                "CAMERA_STARTED", details="Camera started in monitor mode")
            return jsonify({'success': True, 'message': 'Camera started in monitor mode'})
        else:
            return jsonify({'success': False, 'message': 'Failed to start camera. Check connection and permissions.'})
    except Exception as e:
        return jsonify({'success': False, 'message': f'Error starting camera: {str(e)}'})

@api_bp.route('/stop_camera_only', methods=['POST'])
def stop_camera_only():
    """Stop camera but keep system ready"""
    try:
        attendance_system.camera_manager.stop_camera()
        db_manager.log_security_event(
            "CAMERA_STOPPED", details="Camera stopped from monitor mode")
        return jsonify({'success': True, 'message': 'Camera stopped'})
    except Exception as e:
        return jsonify({'success': False, 'message': f'Error stopping camera: {str(e)}'})

# =========================
# AUTHENTICATION ENDPOINTS
# =========================

@api_bp.route('/login', methods=['POST'])
def login():
    """User login endpoint"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'message': 'Invalid request data'})

        username = data.get('username', '').strip()
        password = data.get('password', '')

        if not username or not password:
            return jsonify({'success': False, 'message': 'Username and password are required'})

        user = db_manager.verify_user(username, password)

        if user and user.get('is_active'):
            # Set session variables
            session['user_id'] = user['user_id']
            session['username'] = user['username']
            session['role'] = user['role']
            session['logged_in'] = True

            db_manager.log_security_event(
                "USER_LOGIN",
                details=f"User '{username}' logged in successfully"
            )

            return jsonify({
                'success': True,
                'message': 'Login successful',
                'user': {
                    'username': user['username'],
                    'role': user['role']
                },
                'redirect': '/dashboard'
            })
        else:
            db_manager.log_security_event(
                "FAILED_LOGIN",
                details=f"Failed login attempt for username '{username}'"
            )
            return jsonify({'success': False, 'message': 'Invalid username or password'})

    except Exception as e:
        logger.exception("Login error")
        return jsonify({'success': False, 'message': f'Login error: {str(e)}'})

@api_bp.route('/logout', methods=['POST'])
def logout():
    """User logout endpoint"""
    try:
        username = session.get('username', 'Unknown')
        db_manager.log_security_event(
            "USER_LOGOUT",
            details=f"User '{username}' logged out"
        )

        # Clear session
        session.clear()

        return jsonify({
            'success': True,
            'message': 'Logged out successfully',
            'redirect': '/login'
        })

    except Exception as e:
        logger.exception("Logout error")
        return jsonify({'success': False, 'message': f'Logout error: {str(e)}'})

@api_bp.route('/check_session', methods=['GET'])
def check_session():
    """Check if user is logged in"""
    if 'user_id' in session:
        return jsonify({
            'logged_in': True,
            'user': {
                'username': session.get('username'),
                'role': session.get('role')
            }
        })
    return jsonify({'logged_in': False})

@api_bp.route('/users', methods=['GET'])
@login_required
def get_users():
    """Get all users (admin only)"""
    try:
        users = db_manager.get_all_users()
        return jsonify({'success': True, 'users': users})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@api_bp.route('/users', methods=['POST'])
@login_required
def add_user():
    """Add new user (admin only)"""
    try:
        data = request.get_json()
        username = data.get('username', '').strip()
        password = data.get('password', '')
        email = data.get('email', '')
        role = data.get('role', 'admin')

        if not username or not password:
            return jsonify({'success': False, 'message': 'Username and password are required'})

        success, message = db_manager.add_user(username, password, email, role)

        if success:
            db_manager.log_security_event(
                "USER_CREATED",
                details=f"New user '{username}' created with role '{role}'"
            )

        return jsonify({'success': success, 'message': message})

    except Exception as e:
        logger.exception("Error adding user")
        return jsonify({'success': False, 'message': str(e)})

@api_bp.route('/change_password', methods=['POST'])
@login_required
def change_password():
    """Change user password"""
    try:
        data = request.get_json()
        current_password = data.get('current_password', '')
        new_password = data.get('new_password', '')

        if not current_password or not new_password:
            return jsonify({'success': False, 'message': 'Current and new password are required'})

        # Verify current password
        username = session.get('username')
        user = db_manager.verify_user(username, current_password)

        if not user:
            return jsonify({'success': False, 'message': 'Current password is incorrect'})

        # Update password
        success, message = db_manager.update_user_password(user['user_id'], new_password)

        if success:
            db_manager.log_security_event(
                "PASSWORD_CHANGED",
                details=f"Password changed for user '{username}'"
            )

        return jsonify({'success': success, 'message': message})

    except Exception as e:
        logger.exception("Error changing password")
        return jsonify({'success': False, 'message': str(e)})
